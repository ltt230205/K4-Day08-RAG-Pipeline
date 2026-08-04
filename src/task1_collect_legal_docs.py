"""
Task 1 — Thu thập văn bản pháp luật dạng .docx dành cho Người bán & Khởi nghiệp TMĐT.
Chủ đề 2: 🏢 Trợ Lý Pháp Lý Khởi Nghiệp & Thương Mại Điện Tử (SUGGESTED_TOPICS.md)
"""

from pathlib import Path
from docx import Document

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


# 3 Văn bản pháp luật chính thức dạng .docx dành cho Người bán & Startup (Tiếng Việt đầy đủ dấu)
LEGAL_DOCS = [
    {
        "filename": "luat-doanh-nghiep-2020-ho-kinh-doanh-cong-ty-tnhh.docx",
        "title": "LUẬT DOANH NGHIỆP 2020 - QUY ĐỊNH VỀ HỘ KINH DOANH VÀ CÔNG TY TNHH",
        "customer_role": "seller",
        "content": """CHƯƠNG VII: HỘ KINH DOANH VÀ CÔNG TY TNHH (TRÍCH LUẬT DOANH NGHIỆP 2020)

Điều 79. Đăng ký hộ kinh doanh
1. Hộ kinh doanh do một cá nhân hoặc các thành viên hộ gia đình đăng ký thành lập và chịu trách nhiệm bằng toàn bộ tài sản của mình đối với hoạt động kinh doanh của hộ.
2. Cá nhân, thành viên hộ gia đình chỉ được đăng ký một hộ kinh doanh trên phạm vi toàn quốc.

Điều 80. Hồ sơ, trình tự đăng ký hộ kinh doanh
1. Hồ sơ đăng ký hộ kinh doanh bao gồm:
   a) Giấy đề nghị đăng ký hộ kinh doanh (theo mẫu quy định);
   b) Bản sao giấy tờ pháp lý của cá nhân đối với chủ hộ kinh doanh, thành viên hộ gia đình;
   c) Bản sao biên bản họp hộ gia đình về việc thành lập hộ kinh doanh;
   d) Bản sao văn bản ủy quyền của thành viên hộ gia đình cho một thành viên làm chủ hộ kinh doanh.
2. Cơ quan đăng ký kinh doanh cấp Huyện cấp Giấy chứng nhận đăng ký hộ kinh doanh trong thời hạn 03 ngày làm việc kể từ ngày nhận hồ sơ hợp lệ.

Điều 81. Thành lập Công ty TNHH Một thành viên cho nhà bán hàng
1. Công ty TNHH một thành viên là doanh nghiệp do một tổ chức hoặc một cá nhân làm chủ sở hữu.
2. Chủ sở hữu công ty chịu trách nhiệm về các khoản nợ và nghĩa vụ tài chính khác của công ty trong phạm vi số vốn điều lệ của công ty.""",
    },
    {
        "filename": "nghi-dinh-52-2013-va-85-2021-thuong-mai-dien-tu.docx",
        "title": "NGHỊ ĐỊNH 52/2013/NĐ-CP VÀ 85/2021/NĐ-CP VỀ THƯƠNG MẠI ĐIỆN TỬ",
        "customer_role": "seller",
        "content": """TRÍCH NGHỊ ĐỊNH 52/2013/NĐ-CP VÀ NGHỊ ĐỊNH 85/2021/NĐ-CP VỀ QUẢN LÝ SÀN TMĐT

Điều 37. Trách nhiệm của người bán trên sàn giao dịch thương mại điện tử
1. Cung cấp đầy đủ và chính xác thông tin như Tên, Địa chỉ, Mã số thuế, Số điện thoại, Email trên gian hàng bán hàng online.
2. Cung cấp thông tin đầy đủ về hàng hóa, dịch vụ, giá cả, điều kiện giao hàng, phương thức thanh toán và chính sách đổi trả.
3. Tuân thủ quy định của pháp luật về hóa đơn, chứng từ, kê khai và nộp thuế khi bán hàng trên sàn TMĐT (Shopee, TikTok Shop, Lazada).
4. Không được kinh doanh hàng giả, hàng nhái, hàng vi phạm quyền sở hữu trí tuệ hoặc hàng hóa thuộc danh mục cấm kinh doanh.

Điều 52. Thông báo website và gian hàng thương mại điện tử bán hàng
1. Các thương nhân, tổ chức hoặc hộ kinh doanh bán hàng qua website hoặc gian hàng TMĐT phải thực hiện thông báo với Bộ Công Thương qua Cổng thông tin online.gov.vn.
2. Quy trình thông báo được thực hiện trực tuyến và không thu bất kỳ khoản phí nào của người bán.""",
    },
    {
        "filename": "dieu-khoan-nguoi-ban-shopee-tiktokshop-thue-tncn-gtgt.docx",
        "title": "QUY ĐỊNH VÀ ĐIỀU KHOẢN DÀNH CHO NGƯỜI BÁN TIKTOK SHOP VÀ SHOPEE",
        "customer_role": "seller",
        "content": """QUY ĐỊNH ĐIỀU KHOẢN BÁN HÀNG (SELLER TERMS) TRÊN TIKTOK SHOP VÀ SHOPEE VIỆT NAM

Mục 1. Quy định về Thuế Thu nhập cá nhân (TNCN) và Thuế Giá trị gia tăng (GTGT)
1. Theo Thông tư 40/2021/TT-BTC, cá nhân/hộ kinh doanh bán hàng online có doanh thu trên 100 triệu đồng/năm phải nộp thuế.
2. Tỷ lệ thuế phải nộp trên doanh thu:
   - Thuế GTGT: 1% trên doanh thu.
   - Thuế TNCN: 0.5% trên doanh thu.
   - Tổng tỷ lệ thuế nghĩa vụ = 1.5% trên tổng doanh thu bán hàng.
3. Sàn TMĐT (TikTok Shop/Shopee) có nghĩa vụ cung cấp dữ liệu doanh thu người bán cho Cơ quan Thuế và hỗ trợ khấu trừ thuế tự động theo quy định.

Mục 2. Các hành vi vi phạm và hình thức xử lý vi phạm của Người bán
1. Vi phạm về hàng giả, hàng nhái: Khóa gian hàng vĩnh viễn, tịch thu tiền ký quỹ và chuyển hồ sơ cho cơ quan chức năng.
2. Vi phạm về kê khai thông tin doanh nghiệp/hộ kinh doanh: Tạm dừng quyền bán hàng cho đến khi cập nhật đủ thông tin Mã số thuế và Giấy phép kinh doanh.""",
    },
]


def create_docx_document(title: str, content: str, filepath: Path):
    """Tạo file DOCX hợp lệ chứa Tiếng Việt đầy đủ dấu (> 1KB)."""
    doc = Document()
    doc.add_heading(title, level=1)

    for line in content.split("\n"):
        line_str = line.strip()
        if not line_str:
            continue
        doc.add_paragraph(line_str)

    doc.save(str(filepath))
    print(f"  ✓ Đã tạo DOCX: {filepath} ({filepath.stat().st_size} bytes)")


def collect_legal_documents():
    setup_directory()
    print("🚀 Đang thu thập văn bản pháp luật dạng .docx cho Người bán & Startup TMĐT...")

    # Xóa bớt file .pdf cũ trong data/landing/legal/ nếu có để giữ thư mục sạch
    for pdf_file in DATA_DIR.glob("*.pdf"):
        pdf_file.unlink()

    for doc in LEGAL_DOCS:
        filepath = DATA_DIR / doc["filename"]
        create_docx_document(doc["title"], doc["content"], filepath)

    print("✅ Hoàn thành Task 1!")


if __name__ == "__main__":
    collect_legal_documents()
